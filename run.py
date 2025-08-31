#!/usr/bin/env python3

import json
from docx import Document

# Load resume data from JSON file
with open('resume_data.json', 'r') as f:
    resume_data = json.load(f)

# Create a new Word document
doc = Document()

# Add personal information
doc.add_heading(resume_data['personal_info']['name'], 0)
doc.add_paragraph(resume_data['personal_info']['email'])
doc.add_paragraph(resume_data['personal_info']['phone'])
doc.add_paragraph(resume_data['personal_info']['address'])

# Add experience section
doc.add_heading('Experience', level=1)
for exp in resume_data['experience']:
    doc.add_paragraph(f"{exp['company']} - {exp['position']}")
    doc.add_paragraph(exp['duration'])
    for responsibility in exp['responsibilities']:
        doc.add_paragraph(responsibility, style='List Bullet')

# Add education section
doc.add_heading('Education', level=1)
for edu in resume_data['education']:
    doc.add_paragraph(f"{edu['degree']} - {edu['institution']}")
    doc.add_paragraph(edu['year'])

# Save the document
doc.save('resumes/resume.docx')
